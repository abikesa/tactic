from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
doc = Document()

# Title
title = doc.add_heading('Ukubona Labs: Continuity Solutions for Academic Institutions', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Strategic Teaching, Research, and Workflow Infrastructure for Higher Education')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

# Spacer
doc.add_paragraph('')

# Section: About Ukubona Labs
doc.add_heading('About Ukubona Labs', level=2)
doc.add_paragraph(
    "Ukubona Labs is a specialized academic infrastructure firm founded by Abimereki Muzaale, MD, MPH, PhD(c), a Johns "
    "Hopkins–trained physician-scientist and platform architect. We provide high-value continuity services for institutions "
    "facing hiring freezes, federal funding cuts, and international visa disruptions. Our focus is on epistemic infrastructure: "
    "helping institutions not just survive, but thrive, by building resilient teaching, research, and digital systems."
)

# Section: Our Services
doc.add_heading('Our Services', level=2)
doc.add_paragraph(
    "We deliver rapid-deployment, high-impact support in five key areas:"
)
doc.add_paragraph("1. Teaching Infrastructure – Flipped classrooms, digital curriculum, instructional design.")
doc.add_paragraph("2. Research Continuity – Reproducible workflows, data cleaning, grant compliance, reproducibility protocols.")
doc.add_paragraph("3. Workflow Automation – Research dashboards, personalized risk calculators, platform integration.")
doc.add_paragraph("4. International Continuity – Remote mentoring, visa-interrupted research coverage.")
doc.add_paragraph("5. Academic Infrastructure – E-verified subcontracting for research centers and hospitals.")

# Section: Engagement Options
doc.add_heading('Engagement Options', level=2)
doc.add_paragraph(
    "We offer flexible partnerships tailored to your division's needs:"
)
doc.add_paragraph("• Monthly Retainers: 10–40 hours of continuity support.")
doc.add_paragraph("• Grant Subcontracting: SAM-registered, DUNS-compliant academic support.")
doc.add_paragraph("• Project-Based Delivery: Specific platform builds or instructional design packages.")
doc.add_paragraph("• Teaching & Advising: Available for semester-based engagements or short-term teaching roles via LLC.")

# Section: Our Advantage
doc.add_heading('Our Advantage', level=2)
doc.add_paragraph(
    "Because we operate outside of university payroll systems, Ukubona Labs can help bypass hiring freezes while meeting urgent needs. "
    "We bring Hopkins-grade training, global health fluency, and a unique grasp of platform epistemology to your institutional challenges."
)

# Section: Contact
doc.add_heading('Contact', level=2)
doc.add_paragraph("Abimereki Muzaale, MD, MPH, PhD(c) – Founder & Director")
doc.add_paragraph("Email: abikesa@gmail.com")
doc.add_paragraph("Website: abikesa.github.io")
doc.add_paragraph("SAM.gov Registered | E-Verified | DUNS Available Upon Request")

# Save the document
doc_path = "/mnt/data/Ukubona_Labs_Continuity_OnePager.docx"
doc.save(doc_path)

doc_path
